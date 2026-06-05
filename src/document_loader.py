import os
import logging
from pathlib import Path
from typing import List
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def _load_pdf(file_path: str) -> List[Document]:
    from pypdf import PdfReader
    reader = PdfReader(file_path, strict=False)
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": os.path.basename(file_path), "page": i + 1}
            ))
    logger.info("PDF loaded: %s (%d pages)", file_path, len(docs))
    return docs


def _load_docx(file_path: str) -> List[Document]:
    from docx import Document as DocxDoc
    doc = DocxDoc(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(page_content=text, metadata={"source": os.path.basename(file_path)})]


def _load_text(file_path: str) -> List[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    return [Document(page_content=content, metadata={"source": os.path.basename(file_path)})]


def load_document(file_path: str) -> List[Document]:
    ext = Path(file_path).suffix.lower()
    loaders = {".pdf": _load_pdf, ".docx": _load_docx, ".txt": _load_text, ".md": _load_text}
    loader = loaders.get(ext)
    if not loader:
        raise ValueError(f"Unsupported file type: {ext}")
    return loader(file_path)
